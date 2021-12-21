# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.
'''
使用 \ufeff 分隔字符串
\u200b 转为 1
\u200c 转为 0
'''
import docx
import docx.shared
import urllib.request as ur
'''
pip3 install python-docx
'''


# 使用自定义字符样式
# document.add_paragraph('').add_run('正月里采花无哟花采，二月间采花花哟正开，二月间采花花哟正开。三月里桃花红哟似海，

def string2bin(s):
    return ' '.join([bin(ord(c)).replace('0b', '') for c in s])

def bin2string(s):
    return ''.join([chr(i) for i in [int(b, 2) for b in s.split(' ')]])

def bin2unicode(s):
    str = ""
    for c in s:
        if c == '0':
            str = str + "\u200c\ufeff"
        if c == '1':
            str = str + "\u200b\ufeff"
        if c == ' ':
            str = str + "\ufeff\ufeff"
    return  str

def unicode2bin(s):
    str = ""
    for c, j in zip(s[::2], s[1::2]):
        if c == '\u200c':
            str = str + "0"
        if c == '\u200b':
            str = str + "1"
        if c == '\ufeff':
            str = str + " "

    return  str

'''def unicode2bin(s):
    str = ""
    for c, j in zip(s[::2], s[1::2]):
        if c == '\u200c':
            str = str + "1"
        if c == '\u200b':
            str = str + " "
        if c == '\ufeff':
            str = str + "0"

    return  str'''

def menu():
    print("输入1是加水印，输入2是解出水印")



# 字符串转字节
stra = '中国'
bit = string2bin(stra)
print(bit)#输出6个字节
unicode = bin2unicode(bit)
print(unicode)
bin2 = unicode2bin(unicode)
print(bin2)
str = bin2string(bit)
print(str)#输出6个字节

''' print("请输入想要加的水印")
 str = input()
 bit = string2bin(str)
 unicode = bin2unicode(bit)
 print("请输入水印想加在什么文字后")
 target = input()'''

document=docx.Document("test.docx")
menu()
choice = input()
if choice == "1":
    str1 = "nuaa"
    bit = string2bin(str1)
    unicode = bin2unicode(bit)
    for paragraph in document.paragraphs:
        paragraph.add_run('unicode').font.hidden = True
    document.save('test.docx')
if choice == "2":
    str1 = "nuaa"
    bit = string2bin(str1)
    target = bin2unicode(bit)
    bin2 =unicode2bin(target)
    stree=bin2string(bin2)
    i = 1
    for paragraph in document.paragraphs:
        i = i +1
        for run in paragraph.runs:
            if target in run.text:
                line = str(i)
                print("第" + line + "行存在水印：" + stree)


#输出段落编号及段落内容


'''
from docx import Document
import re

doc = Document(r"D:\论文.docx")
restr = '"(?:[^"])*"'

for p in doc.paragraphs:
    matchRet = re.findall(restr, p.text)
    for r in matchRet:
        p.text = p.text.replace(r, '“' + r[1:-1] + '”')
doc.save(r'D:\论文_修正.docx')
'''



'''def print_hi(name):
    # Use a breakpoint in the code line below to debug your script.
    print(f'Hi, {name}')  # Press Ctrl+F8 to toggle the breakpoint.


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    print_hi('PyCharm')
'''
# See PyCharm help at https://www.jetbrains.com/help/pycharm/
